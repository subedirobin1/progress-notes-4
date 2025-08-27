# app.py
import io
import traceback
from datetime import datetime
import html

import streamlit as st

# Import python-docx
from docx import Document
from docx.shared import Cm, Pt
try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_ALIGN
except Exception:
    WD_ALIGN = None

# ---------------- Instruction Manual ---------------- #

st.set_page_config(page_title="Patient Progress Notes Generator", layout="centered")
st.title("📋 Patient Progress Notes Generator")

st.markdown("""
## 📖 Instruction Manual

This app generates **daily patient progress notes** from a `.docx` file that contains a **table**.  
The table must have **4 columns** with the following headings:

| Column 1 (Patient details) | Column 2 (Issues) | Column 3 (Lab results / On review) | Column 4 (Plan) |
|-----------------------------|-------------------|-------------------------------------|-----------------|
| Name, ID, Age, Ward         | List of issues    | Latest labs, review findings        | Management plan |

### ✅ Steps to Use:
1. Open Microsoft Word (or similar).
2. Create a table with **4 columns** as shown above (you may include a header row — the app will skip it).
3. Enter each patient in a **new row**:
   - **Patient details:** e.g. `Mr. John Smith, ID 12345, Ward 3B`
   - **Issues:** Write one issue per line.
   - **Lab results / On review:** Enter bloods, imaging, or review notes.
   - **Plan:** Enter the treatment or plan.
4. Save the file as **Word Document (.docx)**.
5. Upload the `.docx` file below.
6. Type your **Team and Members** in the box.
7. Preview the notes.
8. Click **Generate Progress Notes** to download a formatted `.docx`.

---
""", unsafe_allow_html=True)

# ---------------- Helper Functions ---------------- #

def get_text_from_cell(cell):
    return "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])

def parse_docx(uploaded_file):
    try:
        data = uploaded_file.getvalue()
    except Exception:
        uploaded_file.seek(0)
        data = uploaded_file.read()
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        st.error("Could not parse the .docx file. Make sure it is valid.")
        st.exception(e)
        return []

    patients = []
    if not doc.tables:
        return patients

    table = doc.tables[0]
    for row_idx, row in enumerate(table.rows):
        if len(row.cells) < 4:
            continue
        pinfo = get_text_from_cell(row.cells[0])
        issues = get_text_from_cell(row.cells[1])
        labs = get_text_from_cell(row.cells[2])
        plan = get_text_from_cell(row.cells[3])

        # Skip header row
        if row_idx == 0:
            header_combined = f"{pinfo} {issues} {labs} {plan}".lower()
            if any(k in header_combined for k in ["patient", "name", "id", "issue", "issues", "lab", "plan", "on review"]):
                continue

        if pinfo.strip():
            patients.append({
                "patient_info": pinfo.strip(),
                "issues": issues.strip(),
                "labs": labs.strip(),
                "plan": plan.strip()
            })
    return patients

def create_progress_notes(patients, team_info):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(5)
    section.right_margin = Cm(2)

    today = datetime.now().strftime("%d %B %Y")

    for idx, p in enumerate(patients):
        para = doc.add_paragraph()
        if WD_ALIGN is not None:
            para.alignment = WD_ALIGN.RIGHT
        run = para.add_run(p["patient_info"])
        run.font.size = Pt(12)
        run.bold = True

        for _ in range(2):
            doc.add_paragraph()

        para = doc.add_paragraph()
        run = para.add_run(team_info if team_info else "")
        run.font.size = Pt(11)
        run.bold = True
        run.italic = True

        doc.add_paragraph()

        para = doc.add_paragraph()
        r1 = para.add_run("Ward Round Notes")
        r1.bold = True
        r1.font.size = Pt(12)
        para.add_run("  " + today).font.size = Pt(11)

        para = doc.add_paragraph()
        para.add_run("Issues").bold = True
        doc.add_paragraph(p["issues"] if p["issues"] else "")

        para = doc.add_paragraph()
        para.add_run("On review").bold = True
        doc.add_paragraph(p["labs"] if p["labs"] else "")

        for _ in range(4):
            doc.add_paragraph()

        para = doc.add_paragraph()
        para.add_run("On examination").bold = True
        for _ in range(4):
            doc.add_paragraph()

        para = doc.add_paragraph()
        para.add_run("Plan").bold = True
        doc.add_paragraph(p["plan"] if p["plan"] else "")

        if idx != len(patients) - 1:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _html_escape_and_br(s):
    if not s:
        return ""
    return html.escape(s).replace("\n", "<br>")

def format_preview_html(p, team_info):
    today = datetime.now().strftime("%d %B %Y")
    issues_html = _html_escape_and_br(p.get("issues", ""))
    labs_html = _html_escape_and_br(p.get("labs", ""))
    plan_html = _html_escape_and_br(p.get("plan", ""))

    html_parts = f"""
<div style="text-align:right; font-size:16px; font-weight:bold;">
{_html_escape_and_br(p.get('patient_info',''))}
</div>
<br><br>
<div style="font-size:14px; font-style:italic; font-weight:bold; color:#444;">
{_html_escape_and_br(team_info or "")}
</div>
<br>
<div style="font-size:16px; font-weight:bold;">
Ward Round Notes <span style="font-size:14px; font-style:italic; font-weight:normal;">{today}</span>
</div>
<div style="font-weight:bold;">Issues</div>
<div>{issues_html}</div>
<div style="font-weight:bold;">On review</div>
<div>{labs_html}</div>
<br><br><br><br>
<div style="font-weight:bold;">On examination</div>
<br><br><br><br>
<div style="font-weight:bold;">Plan</div>
<div>{plan_html}</div>
"""
    return html_parts

def create_sample_docx_bytes():
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Patient details"
    hdr_cells[1].text = "Issues"
    hdr_cells[2].text = "On review"
    hdr_cells[3].text = "Plan"

    rows = [
        ("Mr John Smith, ID 1001, Ward 3B", "Fever\nDyspnoea", "WCC 14.2\nCRP 45", "Start IV antibiotics\nCXR"),
        ("Ms Alice Brown, ID 1002, Ward 4A", "Abdominal pain", "WCC 9.8\nLFTs normal", "CT abdomen\nNPO"),
        ("Mr Bob Lee, ID 1003, Ward 2C", "Post-op Day 2\nPain controlled", "Hb 110", "Analgesia PRN\nPhysio")
    ]
    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = r[0]
        row_cells[1].text = r[1]
        row_cells[2].text = r[2]
        row_cells[3].text = r[3]

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.getvalue()

# ---------------- Streamlit UI ---------------- #

try:
    st.subheader("👥 Team")
    team_info = st.text_input("👨‍⚕️ Team and Members", "")

    st.markdown("**Need an example file?** Download and open this sample to see the required table format.")
    sample_bytes = create_sample_docx_bytes()
    st.download_button("📄 Download sample_patient_list.docx", data=sample_bytes, file_name="sample_patient_list.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    uploaded_file = st.file_uploader("📂 Upload patient list (.docx)", type=["docx"])

    if uploaded_file is not None:
        patients = parse_docx(uploaded_file)
        if not patients:
            st.error("No valid patient rows found in the document (or failed to parse). Try the sample file to check format.")
        else:
            st.success(f"✅ Found {len(patients)} patient(s) in the file.")

            st.subheader("🔎 Live Preview")
            for idx, p in enumerate(patients):
                st.markdown(format_preview_html(p, team_info), unsafe_allow_html=True)
                if idx != len(patients) - 1:
                    st.markdown("<hr style='border: 2px dashed #bbb; margin: 40px 0;'>", unsafe_allow_html=True)

            if st.button("Generate Progress Notes"):
                try:
                    output = create_progress_notes(patients, team_info)
                    st.success("Progress notes generated — click to download:")
                    st.download_button(
                        label="📥 Download Progress Notes (.docx)",
                        data=output.getvalue() if hasattr(output, "getvalue") else output,
                        file_name="progress_notes.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error("An error occurred while creating the DOCX.")
                    st.exception(traceback.format_exc())

    st.markdown("""
<hr style="border:1px solid #bbb; margin-top:50px; margin-bottom:20px;">
<div style="text-align:center; font-size:16px;">
✨ Made with ❤️ & ☕ by <b>Rabindra Subedi</b> ✨
</div>
""", unsafe_allow_html=True)

except Exception:
    st.error("An unexpected error occurred. See details below:")
    st.exception(traceback.format_exc())
